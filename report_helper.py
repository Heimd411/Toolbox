from docx import Document

document = Document()

document.add_heading(input('Document title:\n'), 0)

while True:
    document.add_heading(input('Image heading:\n'), 4)
    document.add_picture(input('Image filename:\n'))
    document.add_paragraph(input('Comment:\n'))
    document.save('test.docx')

